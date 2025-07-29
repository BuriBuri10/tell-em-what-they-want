# core/prompts/report_saver.py

import os
from pathlib import Path
from datetime import datetime
from workflows.state import GraphState
from logs.logging_config import logger

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# core/utils/campaign_recs.py

class CampaignReportSaver:
    @staticmethod
    def save(state: GraphState, user_id: str, query: str, timestamp: str) -> None:
        base_filename = f"{user_id}_{timestamp}"
        txt_filename = f"{base_filename}.txt"
        docx_filename = f"{base_filename}.docx"

        output_dir = Path(__file__).resolve().parent.parent.parent / "campaign_outputs"
        output_dir.mkdir(parents=True, exist_ok=True)

        txt_path = output_dir / txt_filename
        docx_path = output_dir / docx_filename

        report = f"""📊 CAMPAIGN REPORT
───────────────────────
User ID: {user_id}
Timestamp: {timestamp}
Campaign Query: {query}

User Segment: {state.user_segment or 'N/A'}
Campaign Objective: {state.campaign_objective or 'N/A'}
Recommendation: {state.campaign_recommendation or 'N/A'}
Generated Ad Copy: {state.generated_ad or 'N/A'}
Human Feedback: {state.ad_feedback or 'No feedback yet'}

Recommendation:
{state.campaign_recommendation or 'N/A'}

Generated Ad Copy:
{state.generated_ad or 'N/A'}

Human Feedback:
{state.ad_feedback or 'No feedback yet'}
"""
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(report)

        # Save docx
        CampaignReportSaver._save_as_docx(state, user_id, query, docx_path)

        logger.info(f"Campaign report saved: {txt_path}, {docx_path}")

# class CampaignReportSaver:
#     @staticmethod
#     def save(state: GraphState, user_id: str, query: str, timestamp: str) -> None:
#         # timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

#         filename = f"{user_id}_{timestamp}.txt"

#         # Resolve absolute path
#         output_dir = Path(__file__).resolve().parent.parent.parent / "campaign_outputs"
#         output_dir.mkdir(parents=True, exist_ok=True)

#         file_path = output_dir / filename

#         # filename = f"campaign_outputs/{user_id}_{timestamp}.txt"
#         # os.makedirs("campaign_outputs", exist_ok=True)

#         report = f"""📊 CAMPAIGN REPORT
#                 ─────────────────────────
#                 User ID: {user_id}
#                 Timestamp: {timestamp}
#                 Campaign Query: {query}

#                 User Segment: {state.user_segment or 'N/A'}
#                 Campaign Objective: {state.campaign_objective or 'N/A'}
#                 Recommendation: {state.campaign_recommendation or 'N/A'}
#                 Generated Ad Copy: {state.generated_ad or 'N/A'}
#                 Human Feedback: {state.ad_feedback or 'No feedback yet'}

#                 Recommendation:
#                 {state.campaign_recommendation or 'N/A'}

#                 Generated Ad Copy:
#                 {state.generated_ad or 'N/A'}

#                 Human Feedback:
#                 {state.ad_feedback or 'No feedback yet'}
#                 """
#         # with open(filename, "w", encoding="utf-8") as f:
#         with open(file_path, "w", encoding="utf-8") as f:
#             f.write(report)

#         # Save as .docx
#         docx_path = output_dir / f"{filename}.docx"
#         CampaignReportSaver._save_as_docx(state, user_id, query, docx_path)

#         logger.info(f"Campaign report saved: {file_path}, {docx_path}")

    @staticmethod
    def _save_as_docx(state: GraphState, user_id: str, query: str, docx_path: Path) -> None:
        doc = Document()

        # Title
        title = doc.add_heading("📊 CAMPAIGN REPORT", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Meta info
        doc.add_paragraph(f"User ID: {user_id}")
        doc.add_paragraph(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Campaign Query: {query}")

        doc.add_paragraph("─" * 30)

        # Campaign details
        doc.add_paragraph(f"User Segment: {state.user_segment or 'N/A'}")
        doc.add_paragraph(f"Campaign Objective: {state.campaign_objective or 'N/A'}")
        doc.add_paragraph(f"Recommendation: {state.campaign_recommendation or 'N/A'}")
        doc.add_paragraph(f"Generated Ad Copy: {state.generated_ad or 'N/A'}")
        doc.add_paragraph(f"Human Feedback: {state.ad_feedback or 'No feedback yet'}")

        # Save the .docx file
        doc.save(docx_path)
